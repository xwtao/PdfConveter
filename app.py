from flask import Flask, request, jsonify, send_file
from werkzeug.utils import secure_filename
import os
from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import io

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max-limit

# 确保上传目录存在
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() == 'pdf'

@app.route('/health')
def health_check():
    return jsonify({"status": "healthy"}), 200

@app.route('/convert', methods=['POST'])
def convert_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed'}), 400

    try:
        # 保存上传的文件
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(pdf_path)

        # 创建Word文档
        doc = Document()
        pdf = PdfReader(pdf_path)

        # 转换每一页
        for page in pdf.pages:
            text = page.extract_text()
            doc.add_paragraph(text)

            # 提取图片
            for image in page.images:
                image_data = image.data
                image_stream = io.BytesIO(image_data)
                img = Image.open(image_stream)
                
                # 保存图片到临时文件
                img_path = os.path.join(app.config['UPLOAD_FOLDER'], f"temp_{image.name}")
                img.save(img_path)
                
                # 添加图片到文档
                doc.add_picture(img_path)
                os.remove(img_path)  # 删除临时图片文件

        # 保存Word文档
        docx_filename = os.path.splitext(filename)[0] + '.docx'
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        doc.save(docx_path)

        # 清理PDF文件
        os.remove(pdf_path)

        # 返回转换后的文件
        return send_file(
            docx_path,
            as_attachment=True,
            download_name=docx_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True) 